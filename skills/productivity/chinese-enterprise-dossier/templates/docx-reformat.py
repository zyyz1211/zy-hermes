# -*- coding: utf-8 -*-
"""
公文格式重排版脚本 — 对已有 .docx 按 GB/T 9704-2012 标准重排

使用方法：
  1. 修改 INPUT / OUTPUT 路径
  2. 调整段落分类规则中 'title' 匹配的文档标题文本
  3. ~/.local/open-webui-venv/bin/python3.11 script.py
"""
import os, re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ========== 配置 ==========
INPUT = '/mnt/c/Users/<user>/Desktop/01_客户管理/0.客户档案/某石化企业/某石化企业档案.docx'
OUTPUT = '/mnt/c/Users/<user>/Desktop/01_客户管理/0.客户档案/某石化企业/某石化企业档案_公文版.docx'

# 文档标题文本（精确匹配，用于分类为 'title'）
DOC_TITLE_TEXT = '某石化企业档案'


# ========== 格式工具函数 ==========

def set_page_layout(doc):
    """A4 + 公文页边距（上37/下35/左28/右26mm）"""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)


def set_run_font(run, font_name, size_pt=16, bold=False):
    """统一设置 run 字体"""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    rPr = run._r.get_or_add_rPr()
    for child in list(rPr):
        if child.tag == qn('w:rFonts'):
            rPr.remove(child)
    rFonts = rPr.makeelement(qn('w:rFonts'), {
        qn('w:ascii'): font_name,
        qn('w:hAnsi'): font_name,
        qn('w:eastAsia'): font_name,
    })
    rPr.append(rFonts)


def set_paragraph_format(p, indent=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_pt=28):
    """统一段落格式：固定行距 + 缩进 + 两端对齐"""
    pPr = p._p.get_or_add_pPr()
    # 清除旧 spacing
    for child in list(pPr):
        if child.tag == qn('w:spacing'):
            pPr.remove(child)
    # 固定行距（28磅 = 560 单位）
    spacing = pPr.makeelement(qn('w:spacing'), {
        qn('w:line'): str(line_pt * 20),
        qn('w:lineRule'): 'exact',
    })
    pPr.append(spacing)
    # 清除旧缩进
    for child in list(pPr):
        if child.tag == qn('w:ind'):
            pPr.remove(child)
    if indent:
        pPr.append(pPr.makeelement(qn('w:ind'), {qn('w:firstLine'): '640'}))
    p.alignment = alignment


def clear_run_format(run):
    """清空 run 的所有格式"""
    run.font.name = ''
    run.font.size = None
    run.bold = False
    rPr = run._r.get_or_add_rPr()
    for child in list(rPr):
        if child.tag in (qn('w:rFonts'), qn('w:b'), qn('w:sz'), qn('w:szCs')):
            rPr.remove(child)


def format_paragraph(para, fmt_type):
    """按类型格式化段落"""
    if fmt_type == 'title':
        # 文档标题：黑体二号居中
        for run in para.runs:
            clear_run_format(run)
            set_run_font(run, '黑体', 22, bold=True)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 标题设回默认行距（不设固定行距）
        pPr = para._p.get_or_add_pPr()
        for child in list(pPr):
            if child.tag == qn('w:spacing'):
                pPr.remove(child)

    elif fmt_type == 'info':
        # 基本信息行：仿宋_GB2312 三号加粗，不计缩进
        for run in para.runs:
            clear_run_format(run)
            set_run_font(run, '仿宋_GB2312', 16, bold=True)
        set_paragraph_format(para, indent=False)

    elif fmt_type == 'h1':
        # 一级标题：黑体三号
        for run in para.runs:
            clear_run_format(run)
            set_run_font(run, '黑体', 16, bold=True)
        set_paragraph_format(para)

    elif fmt_type == 'h2':
        # 二级标题：楷体_GB2312 三号
        for run in para.runs:
            clear_run_format(run)
            set_run_font(run, '楷体_GB2312', 16, bold=False)
        set_paragraph_format(para)

    elif fmt_type == 'h3':
        # 三级标题：仿宋_GB2312 加粗 三号
        for run in para.runs:
            clear_run_format(run)
            set_run_font(run, '仿宋_GB2312', 16, bold=True)
        set_paragraph_format(para)

    elif fmt_type == 'body':
        # 正文：仿宋_GB2312 三号
        for run in para.runs:
            clear_run_format(run)
            set_run_font(run, '仿宋_GB2312', 16, bold=False)
        set_paragraph_format(para)

    elif fmt_type == 'table_caption':
        # 表注/说明行：仿宋_GB2312 三号，不计缩进或常规缩进
        for run in para.runs:
            clear_run_format(run)
            set_run_font(run, '仿宋_GB2312', 16, bold=False)
        set_paragraph_format(para)


def set_table_style(table):
    """统一表格格式：宋体五号，表头加粗居中"""
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = (ri == 0)  # 表头加粗
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                    rPr = run._r.get_or_add_rPr()
                    for child in list(rPr):
                        if child.tag == qn('w:rFonts'):
                            rPr.remove(child)
                    rFonts = rPr.makeelement(qn('w:rFonts'), {
                        qn('w:ascii'): '宋体',
                        qn('w:hAnsi'): '宋体',
                        qn('w:eastAsia'): '宋体',
                    })
                    rPr.append(rFonts)


# ========== 段落分类规则 ==========

def classify_paragraph(text):
    """根据文本内容判断段落类型"""
    if not text.strip():
        return None
    if text == DOC_TITLE_TEXT:
        return 'title'
    if text.startswith('客户全称') or text.startswith('客户管家') or text.startswith('建档日期'):
        return 'info'
    if re.match(r'^[一二三四五六七八九十]、', text):
        return 'h1'
    if re.match(r'^（[一二三四五六七八九十]+）', text):
        return 'h2'
    if text.startswith('附表') or text.startswith('附图'):
        return 'table_caption'
    if text.startswith('【'):
        return 'h3'
    if re.match(r'^\d+[\.\、．]?\s', text):
        return 'h3'
    if re.match(r'^（\d+）', text):
        return 'h3'
    # 带冒号的企业基本信息行（如"行业地位：XXX"）
    if re.match(r'^(行业地位|核心盈利板块|支付情况|资信情况|主体信用评级)', text):
        return 'body'
    return 'body'


# ========== 主流程 ==========

doc = Document(INPUT)

# 1. 页面布局
set_page_layout(doc)

# 2. 设置默认样式
style = doc.styles['Normal']
style.font.name = '仿宋_GB2312'
style.font.size = Pt(16)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

# 3. 逐段落格式化
for p in doc.paragraphs:
    text = p.text.strip()
    fmt_type = classify_paragraph(text)
    if fmt_type:
        format_paragraph(p, fmt_type)

# 4. 表格格式化
for table in doc.tables:
    set_table_style(table)

# 5. 保存
os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
doc.save(OUTPUT)
print(f'✅ 已保存：{OUTPUT}')
