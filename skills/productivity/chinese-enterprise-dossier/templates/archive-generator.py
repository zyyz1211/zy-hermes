"""
客户档案生成模板 — 党政机关公文格式 (GB/T 9704-2012)

格式标准：
  - 页面：A4，上37mm/下35mm/左28mm/右26mm
  - 一级标题（一、二、三）：黑体 三号(16pt)
  - 二级标题（（一）（二）（三））：楷体_GB2312 三号(16pt)
  - 三级标题（1. 2. 3.）：仿宋_GB2312 加粗 三号(16pt)
  - 正文：仿宋_GB2312 三号(16pt)，首行缩进2字符，固定行距28磅
  - 表格：宋体 五号(10.5pt)，表头加粗

使用方法：
  1. 修改 _COMPANY / _SHORT / _DATE / _OUTPUT 配置
  2. 填充各章节内容
  3. python3.11 script.py
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ========== 配置（每次使用前修改）==========
_COMPANY = "XX科技集团有限公司"
_SHORT = "XX集团"
_DATE = "2026年5月18日"
_OUTPUT = f"/mnt/c/Users/<user>/Desktop/客户管理工作/客户档案/{_SHORT}档案.docx"


# ========== 格式工具函数 ==========

def set_page_layout(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)


def set_run_font(run, font_name, size_pt, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    rPr = run._r.get_or_add_rPr()
    for child in list(rPr):
        if child.tag == qn('w:rFonts'):
            rPr.remove(child)
    rFonts = rPr.makeelement(qn('w:rFonts'), {
        qn('w:ascii'): font_name, qn('w:hAnsi'): font_name, qn('w:eastAsia'): font_name,
    })
    rPr.append(rFonts)


def set_para_spacing(p, indent=True):
    pPr = p._p.get_or_add_pPr()
    for child in list(pPr):
        if child.tag == qn('w:spacing'): pPr.remove(child)
    spacing = pPr.makeelement(qn('w:spacing'), {
        qn('w:line'): '560', qn('w:lineRule'): 'exact',
    })
    pPr.append(spacing)
    for child in list(pPr):
        if child.tag == qn('w:ind'): pPr.remove(child)
    if indent:
        pPr.append(pPr.makeelement(qn('w:ind'), {qn('w:firstLine'): '640'}))
    # 两端对齐
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_h1(doc, text):
    """一级标题：黑体三号"""
    p = doc.add_paragraph()
    r = p.add_run(text); set_run_font(r, '黑体', 16, bold=True)
    set_para_spacing(p); return p


def add_h2(doc, text):
    """二级标题：楷体_GB2312三号"""
    p = doc.add_paragraph()
    r = p.add_run(text); set_run_font(r, '楷体_GB2312', 16)
    set_para_spacing(p); return p


def add_h3(doc, text):
    """三级标题：仿宋_GB2312加粗三号"""
    p = doc.add_paragraph()
    r = p.add_run(text); set_run_font(r, '仿宋_GB2312', 16, bold=True)
    set_para_spacing(p); return p


def add_body(doc, text, bold_prefix=None, indent=True):
    """正文段落：仿宋_GB2312三号"""
    p = doc.add_paragraph()
    if bold_prefix:
        rb = p.add_run(bold_prefix); set_run_font(rb, '仿宋_GB2312', 16, bold=True)
        rt = p.add_run(text); set_run_font(rt, '仿宋_GB2312', 16)
    else:
        r = p.add_run(text); set_run_font(r, '仿宋_GB2312', 16)
    set_para_spacing(p, indent); return p


def add_info(doc, label, value):
    return add_body(doc, value, bold_prefix=label, indent=False)


def set_table_font(table, font_name='宋体', font_size=Pt(10.5)):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = font_name
                    run.font.size = font_size
                    rPr = run._r.get_or_add_rPr()
                    for child in list(rPr):
                        if child.tag == qn('w:rFonts'): rPr.remove(child)
                    rFonts = rPr.makeelement(qn('w:rFonts'), {
                        qn('w:ascii'): font_name, qn('w:hAnsi'): font_name, qn('w:eastAsia'): font_name,
                    })
                    rPr.append(rFonts)


def make_table_header_bold(table):
    """表头行加粗"""
    for cell in table.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True


# ========== 文档生成示例 ==========

doc = Document()
set_page_layout(doc)

style = doc.styles['Normal']
style.font.name = '仿宋_GB2312'
style.font.size = Pt(16)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

# 文档标题
doc_title = doc.add_heading(f'{_SHORT}档案', level=1)
doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in doc_title.runs:
    set_run_font(run, '黑体', 22)

add_info(doc, '客户全称：', _COMPANY)
add_info(doc, '客户管家：', 'XX')
add_info(doc, '建档日期：', _DATE)
doc.add_paragraph()

# 一、客户基础信息
add_h1(doc, '一、客户基础信息')
add_h2(doc, '（一）主体信息')

table = doc.add_table(rows=9, cols=2, style='Table Grid')
table.cell(0, 0).text = '信息类别'
table.cell(0, 1).text = '填写内容'
make_table_header_bold(table)
rows_data = [
    ('客户全称', _COMPANY), ('客户简称', _SHORT),
    ('法定代表人/实际控制人', ''), ('总部地址/核心经营区域', ''),
    ('企业类型/所属行业', ''), ('核心经营范围/市场地位', ''),
    ('核心竞争力', ''), ('与某工程集团合作主要下属企业', ''),
]
for i, (k, v) in enumerate(rows_data):
    table.cell(i + 1, 0).text = k
    table.cell(i + 1, 1).text = v
set_table_font(table)

# 后续章节同理...
add_h1(doc, '二、发展规划和投资布局')
add_h1(doc, '三、既往合作情况')
add_h1(doc, '四、客户关系维护与管理计划')
add_h1(doc, '五、经营状态动态监控')

os.makedirs(os.path.dirname(_OUTPUT), exist_ok=True)
doc.save(_OUTPUT)
print(f'Saved: {_OUTPUT}')


"""
========== 标题层级速查 ==========

文档标题 'XX集团档案'    -> doc.add_heading(level=1) + set_run_font(黑体, 22)
一级标题 '一、XXX'       -> add_h1(doc, '一、XXX')       # 黑体三号
二级标题 '（一）XXX'     -> add_h2(doc, '（一）XXX')     # 楷体_GB2312三号
三级标题 '1. XXX'        -> add_h3(doc, '1. XXX')        # 仿宋_GB2312加粗三号
正文段落                 -> add_body(doc, '内容')          # 仿宋_GB2312三号
基本信息行               -> add_info(doc, '标签：', '值')
表格内容                 -> add_table(...) + set_table_font(table)  # 宋体五号
表头加粗                 -> make_table_header_bold(table)
"""
