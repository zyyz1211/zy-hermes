# -*- coding: utf-8 -*-
"""政策研究报告生成模板 — 党政机关公文格式 (GB/T 9704-2012)"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_run_font(run, fn, sz, bold=False):
    run.font.name = fn; run.font.size = Pt(sz); run.bold = bold
    rPr = run._r.get_or_add_rPr()
    for c in list(rPr):
        if c.tag == qn('w:rFonts'): rPr.remove(c)
    rf = rPr.makeelement(qn('w:rFonts'), {qn('w:ascii'):fn, qn('w:hAnsi'):fn, qn('w:eastAsia'):fn})
    rPr.append(rf)

def set_sp(p, indent=True):
    pPr = p._p.get_or_add_pPr()
    for c in list(pPr):
        if c.tag in (qn('w:spacing'), qn('w:ind')): pPr.remove(c)
    s = pPr.makeelement(qn('w:spacing'), {qn('w:line'):'560', qn('w:lineRule'):'exact'})
    pPr.append(s)
    if indent: pPr.append(pPr.makeelement(qn('w:ind'), {qn('w:firstLine'):'640'}))
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_title(doc, text):
    """文档标题：黑体二号居中"""
    dt = doc.add_heading(text, level=1); dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in dt.runs: set_run_font(r, '黑体', 22)
    return dt

def add_h1(doc, text):
    """章标题：黑体三号"""
    p = doc.add_paragraph(); r = p.add_run(text); set_run_font(r, '黑体', 16, True); set_sp(p); return p

def add_h2(doc, text):
    """节标题：楷体_GB2312三号"""
    p = doc.add_paragraph(); r = p.add_run(text); set_run_font(r, '楷体_GB2312', 16); set_sp(p); return p

def add_body(doc, text, bold_prefix=None, indent=True):
    """正文：仿宋_GB2312三号"""
    p = doc.add_paragraph()
    if bold_prefix:
        rb = p.add_run(bold_prefix); set_run_font(rb, '仿宋_GB2312', 16, True)
        rt = p.add_run(text); set_run_font(rt, '仿宋_GB2312', 16)
    else:
        r = p.add_run(text); set_run_font(r, '仿宋_GB2312', 16)
    set_sp(p, indent); return p

def add_info(doc, label, value):
    """基本信息行（不缩进）"""
    return add_body(doc, value, bold_prefix=label, indent=False)

def init_doc():
    """初始化文档（页面设置+默认字体）"""
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(3.7); sec.bottom_margin = Cm(3.5)
    sec.left_margin = Cm(2.8); sec.right_margin = Cm(2.6)
    sty = doc.styles['Normal']
    sty.font.name = '仿宋_GB2312'; sty.font.size = Pt(16)
    sty.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    return doc

# ===== 使用示例 =====
if __name__ == '__main__':
    doc = init_doc()
    add_title(doc, 'XX政策研究报告')
    add_info(doc, '', '经营管理部')
    doc.add_paragraph()
    add_h1(doc, '一、政策背景')
    add_body(doc, '此处为正文内容，首行缩进2字符，仿宋_GB2312三号...')
    add_h2(doc, '（一）政策要点')
    add_body(doc, '具体政策解析...')
    add_h1(doc, '二、适配性分析')
    add_body(doc, '结合集团实际分析...')
    add_h1(doc, '三、结论与建议')
    add_body(doc, '1. 第一条建议...', bold_prefix='一是')
    doc.add_paragraph()
    add_info(doc, '', '集团公司经营管理部')
    add_info(doc, '', '2026年X月X日')
    doc.save('/mnt/c/Users/<user>/Desktop/报告.docx')
    print('Saved.')
