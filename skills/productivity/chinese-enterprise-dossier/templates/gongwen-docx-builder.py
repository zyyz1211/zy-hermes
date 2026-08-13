#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
公文格式 Word 文档构建器 — 可复用模块

提供一套标准化的函数，用于从结构化内容（MD 文本、研究笔记等）
生成符合 GB/T 9704-2012 标准的 公文格式 .docx 文档。

用法:
    from templates.gongwen_docx_builder import GongwenDocxBuilder

    builder = GongwenDocxBuilder()
    builder.add_title('调研报告标题')
    builder.add_h1('一、章节标题')
    builder.add_body('正文内容……')
    builder.add_table(['列1','列2'], [['行1','数据'],['行2','数据']])
    builder.save('输出路径.docx')

字体规范:
    标题: 黑体 22pt bold 居中
    H1:   黑体 16pt bold (一、二、三...)
    H2:   楷体 15pt bold (（一）（二）...)
    正文: 仿宋_GB2312 16pt (三号)
    表格: 表头黑体11pt，正文仿宋10pt

页面: A4, 上3.7cm/下3.5cm/左2.8cm/右2.6cm
行距: 1.5倍, 首行缩进32pt
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


class GongwenDocxBuilder:
    """公文格式文档构建器"""

    def __init__(self):
        self.doc = Document()
        self._setup_page()
        self._setup_style()

    def _setup_page(self):
        """设置页面边距（GB/T 9704-2012 标准）"""
        for sec in self.doc.sections:
            sec.top_margin = Cm(3.7)
            sec.bottom_margin = Cm(3.5)
            sec.left_margin = Cm(2.8)
            sec.right_margin = Cm(2.6)

    def _setup_style(self):
        """设置默认正文样式"""
        sty = self.doc.styles['Normal']
        sty.font.name = '仿宋_GB2312'
        sty.font.size = Pt(16)
        sty.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        sty.paragraph_format.line_spacing = 1.5
        sty.paragraph_format.space_after = Pt(0)

    def _add_run(self, paragraph, text, bold=False, size=16, font='仿宋_GB2312', color=None):
        """在段落中添加带格式的文本片段"""
        r = paragraph.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.name = font
        r._element.rPr.rFonts.set(qn('w:eastAsia'), font)
        if color:
            r.font.color.rgb = color
        return r

    def add_title(self, text):
        """添加文档标题（黑体22pt居中加粗）"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(20)
        self._add_run(p, text, bold=True, size=22, font='黑体')

    def add_h1(self, text):
        """添加一级标题（黑体16pt加粗）"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        self._add_run(p, text, bold=True, size=16, font='黑体')

    def add_h2(self, text):
        """添加二级标题（楷体15pt加粗）"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        self._add_run(p, text, bold=True, size=15, font='楷体')

    def add_body(self, text, indent=True):
        """添加正文段落（仿宋16pt，默认首行缩进2字符）"""
        p = self.doc.add_paragraph()
        if indent:
            p.paragraph_format.first_line_indent = Pt(32)
        self._add_run(p, text, size=16)

    def add_body_bold_prefix(self, prefix, text, indent=True):
        """添加正文段落，其中前缀加粗（用于标注性文字如'解读：'）"""
        p = self.doc.add_paragraph()
        if indent:
            p.paragraph_format.first_line_indent = Pt(32)
        self._add_run(p, prefix, bold=True, size=16)
        self._add_run(p, text, size=16)

    def add_quote(self, text, size=14):
        """添加引文（楷体14pt灰色，用于引用原文）"""
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(32)
        r = self._add_run(p, text, size=size, font='楷体')
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    def add_note(self, text, size=14):
        """添加注释说明（楷体14pt灰色）"""
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(32)
        r = self._add_run(p, text, size=size, font='楷体')
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    def add_table(self, headers, data):
        """
        添加表格
        headers: 表头列表，如 ['序号', '名称']
        data: 数据列表，如 [['1', '某某企业'], ['2', '某某项目']]
        """
        t = self.doc.add_table(rows=len(data) + 1, cols=len(headers))
        t.style = 'Table Grid'

        # 表头行
        for j, h in enumerate(headers):
            c = t.rows[0].cells[j]
            c.text = ''
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._add_run(p, h, bold=True, size=11, font='黑体')

        # 数据行
        for ri, row in enumerate(data):
            for ci, txt in enumerate(row):
                c = t.rows[ri + 1].cells[ci]
                c.text = ''
                self._add_run(c.paragraphs[0], str(txt), size=10)

        return t

    def add_blank(self):
        """添加一个空行"""
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def save(self, path):
        """保存文档到指定路径"""
        os.makedirs(os.path.dirname(path), exist_ok=True)
        self.doc.save(path)
        return path
