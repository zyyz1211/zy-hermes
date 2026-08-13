#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Reformat an existing client archive .docx to GB/T 9704-2012 government document format.

Usage:
  ~/.local/open-webui-venv/bin/python3.11 reformat-existing-archive.py

Config:
  Edit INPUT / OUTPUT paths at the top of this script.

Formatting applied:
  - Page layout: A4, margins 上37/下35/左28/右26mm
  - Document title: 黑体 二号(22pt) 居中
  - Info rows (客户全称/管家/建档日期): 仿宋_GB2312 三号(16pt) 加粗, 不计缩进
  - H1 (一、二、三...): 黑体 三号(16pt) 加粗, 首行缩进2字符
  - H2 (（一）（二）...): 楷体_GB2312 三号(16pt), 首行缩进2字符
  - H3 (1. 2. 3. / 数字序号开头): 仿宋_GB2312 三号(16pt) 加粗, 首行缩进2字符
  - Body: 仿宋_GB2312 三号(16pt), 首行缩进2字符, 固定行距28磅, 两端对齐
  - Tables: 宋体 五号(10.5pt), 表头加粗居中, 全边框
"""

import os
import re

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ====== CONFIG ======
INPUT = '/mnt/c/Users/<user>/Desktop/01_客户管理/0.客户档案/某石化企业/某石化企业档案.docx'
OUTPUT = INPUT.replace('.docx', '_公文版.docx')
# ====================


def set_page_layout(doc):
    s = doc.sections[0]
    s.page_width = Cm(21.0)
    s.page_height = Cm(29.7)
    s.top_margin = Cm(3.7)
    s.bottom_margin = Cm(3.5)
    s.left_margin = Cm(2.8)
    s.right_margin = Cm(2.6)


def set_run_font(run, font_name, size_pt=16, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    rPr = run._r.get_or_add_rPr()
    for child in list(rPr):
        if child.tag == qn('w:rFonts'):
            rPr.remove(child)
    rFonts = rPr.makeelement(qn('w:rFonts'), {
        qn('w:ascii'): font_name,
        qn('w:hAnsi'): font_name,
        qn('w:eastAsia'): font_name,
    })
    rPr.append(rFonts)


def set_paragraph_format(p, indent=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_pt=28):
    pPr = p._p.get_or_add_pPr()
    for child in list(pPr):
        if child.tag == qn('w:spacing'):
            pPr.remove(child)
    spacing = pPr.makeelement(qn('w:spacing'), {
        qn('w:line'): str(line_pt * 20),
        qn('w:lineRule'): 'exact',
    })
    pPr.append(spacing)
    for child in list(pPr):
        if child.tag == qn('w:ind'):
            pPr.remove(child)
    if indent:
        pPr.append(pPr.makeelement(qn('w:ind'), {
            qn('w:firstLine'): '640',
        }))
    p.alignment = alignment


def clear_run_format(run):
    run.font.name = ''
    run.font.size = None
    run.bold = False
    rPr = run._r.get_or_add_rPr()
    for child in list(rPr):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('rFonts', 'b', 'sz', 'szCs'):
            rPr.remove(child)


def format_paragraph(para, fmt_type):
    text = para.text.strip()
    for run in para.runs:
        clear_run_format(run)

    if fmt_type == 'title':
        for run in para.runs:
            set_run_font(run, '黑体', 22, bold=True)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    elif fmt_type == 'info':
        for run in para.runs:
            set_run_font(run, '仿宋_GB2312', 16, bold=True)
        set_paragraph_format(para, indent=False)

    elif fmt_type == 'h1':
        for run in para.runs:
            set_run_font(run, '黑体', 16, bold=True)
        set_paragraph_format(para)

    elif fmt_type == 'h2':
        for run in para.runs:
            set_run_font(run, '楷体_GB2312', 16, bold=False)
        set_paragraph_format(para)

    elif fmt_type == 'h3':
        for run in para.runs:
            set_run_font(run, '仿宋_GB2312', 16, bold=True)
        set_paragraph_format(para)

    elif fmt_type == 'body':
        for run in para.runs:
            set_run_font(run, '仿宋_GB2312', 16, bold=False)
        set_paragraph_format(para)

    elif fmt_type == 'table_caption':
        for run in para.runs:
            set_run_font(run, '仿宋_GB2312', 16, bold=False)
        set_paragraph_format(para)


def classify_paragraph(text):
    """Classify a paragraph text into a format type."""
    if not text:
        return None
    if text in ('某石化企业档案',) or re.match(r'^.+集团档案$', text):
        return 'title'
    if text.startswith('客户全称') or text.startswith('客户管家') or text.startswith('建档日期'):
        return 'info'
    if re.match(r'^[一二三四五六七八九十]、', text):
        return 'h1'
    if re.match(r'^（[一二三四五六七八九十]+）', text):
        return 'h2'
    if text.startswith('附表') or text.startswith('附图'):
        return 'table_caption'
    if re.match(r'^\d+[\.\、．\s]', text) or re.match(r'^（\d+）', text):
        return 'h3'
    if text.startswith('【'):
        return 'h3'
    return 'body'


def set_table_style(table):
    for cell in table.rows[0].cells:
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
    for row in table.rows[1:]:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = False
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)


def main():
    doc = Document(INPUT)
    set_page_layout(doc)

    style = doc.styles['Normal']
    style.font.name = '仿宋_GB2312'
    style.font.size = Pt(16)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    for p in doc.paragraphs:
        text = p.text.strip()
        fmt = classify_paragraph(text)
        if fmt:
            format_paragraph(p, fmt)

    for table in doc.tables:
        set_table_style(table)

    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    print(f'Saved: {OUTPUT}')


if __name__ == '__main__':
    main()
