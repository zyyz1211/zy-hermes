# -*- coding: utf-8 -*-
"""公文格式 docx 生成器骨架：方正小标宋标题 / 黑体一级 / 楷体二级 / 仿宋正文。
用法: 复制本文件，替换 add_title/add_h1/add_h2/add_body 之间的正文内容即可。
注意: 本文件顶部 docstring 勿写反斜杠 Windows 路径（会 SyntaxWarning），用正斜杠。
"""
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = r"C:\Users\<user>\Desktop\04_汇报材料\output.docx"  # 改成实际输出路径

doc = Document()

# A4 公文页边距（GB/T 9704-2012 风格：上下 3.7/3.5，左右 2.8/2.6；宽松版 3.0/2.6/2.8/2.6）
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.top_margin, sec.bottom_margin = Cm(3.0), Cm(2.6)
sec.left_margin, sec.right_margin = Cm(2.8), Cm(2.6)


def set_font(run, size=14, bold=False, name='仿宋_GB2312'):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, size=22, bold=True, name='方正小标宋简体')
    p.paragraph_format.space_after = Pt(12)


def add_h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=16, bold=True, name='黑体')
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def add_h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=14, bold=True, name='楷体_GB2312')
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)


def add_body(text, indent=True, size=14):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    r = p.add_run(text)
    set_font(r, size=size)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_note(text):
    """灰色小字注释，用于出处/说明。"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=10.5, name='楷体_GB2312')
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    p.paragraph_format.line_spacing = 1.3


# ============ 正文区：替换以下内容 ============
add_title('关于……的建议')
add_h1('一、总体判断')
add_body('……（来源：莫鼎革同志2026年经营工作会议讲话）')
add_h1('二、重点方向建议')
add_h2('（一）……')
add_body('……')
add_h1('三、风险防控与保障措施')
add_body('……')
# =============================================

doc.save(OUT)
print("saved:", OUT, os.path.getsize(OUT), "bytes")
