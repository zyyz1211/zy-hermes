# -*- coding: utf-8 -*-
"""
从原始文本内容生成格式化的十五五规划纲要PDF
"""
import os
import sys
import re
import subprocess
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def clean_text(text):
    """清理文本中的无关字符"""
    text = re.sub(r'<[^>]+>', '', text)
    text = re.sub(r'  +', ' ', text)
    lines = text.split('\n')
    cleaned = []
    skip_patterns = [
        '您的位置', '当前位置', '首页', '字号', '分享至',
        '附件：', '相关稿件', '主办：', '承办：', '备案',
        '网站标识', '无障碍', '登录', '注册', '扫一扫',
        '微信扫一扫', '返回顶部', '打印本页', '关闭本页',
        '联系我们', '法律声明', '关于本站', '网站地图',
        'Copyright', 'All Rights Reserved',
    ]
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if any(skip in line for skip in skip_patterns):
            continue
        cleaned.append(line)
    return '\n'.join(cleaned)

def create_pdf(province_name, content_text, output_dir):
    """创建格式化DOCX并转换为PDF"""
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    title_text = f"{province_name}国民经济和社会发展\n第十五个五年规划纲要"
    for i, line in enumerate(title_text.split('\n')):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(22) if i == 0 else Pt(18)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        p.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 40)
    run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
    run.font.size = Pt(8)
    p.paragraph_format.space_after = Pt(12)
    
    text = clean_text(content_text)
    lines = text.split('\n')
    
    header_count = 0
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        if line.startswith('**') and line.endswith('**'):
            t = line.replace('*', '').strip()
            if t:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(t)
                run.bold = True
                run.font.size = Pt(14)
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                header_count += 1
                continue
        
        chapter_match = re.match(r'^(第[一二三四五六七八九十百零]+[篇章]|第[一二三四五六七八九十百零]+[节]|[一二三四五六七八九十]+[、．\.]\s*|●\s*)', line)
        if chapter_match or (line.startswith('第') and ('篇' in line or '章' in line)):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(13)
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x66)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            header_count += 1
            continue
        
        t = line.replace('\u2003', '').replace('\u2002', '').replace('□', '').replace('▪', '')
        if len(t) > 5 or header_count > 0:
            p = doc.add_paragraph()
            run = p.add_run(t)
            run.font.size = Pt(10.5)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p.paragraph_format.line_spacing = Pt(22)
            p.paragraph_format.space_after = Pt(3)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f'\n来源：{province_name}人民政府')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p.paragraph_format.space_before = Pt(20)
    
    docx_path = os.path.join(output_dir, f'{province_name}十五五规划纲要.docx')
    doc.save(docx_path)
    print(f"DOCX saved: {docx_path}")
    
    libreoffice = r"C:\Program Files\LibreOffice\program\soffice.exe"
    cmd = [libreoffice, "--headless", "--convert-to", "pdf", docx_path, "--outdir", output_dir]
    result = subprocess.run(cmd, capture_output=True, text=True).returncode
    if result == 0:
        pdf_path = os.path.join(output_dir, f'{province_name}十五五规划纲要.pdf')
        print(f"PDF saved: {pdf_path}")
        os.remove(docx_path)
    else:
        print(f"PDF conversion failed for {province_name}")
    
    return pdf_path

if __name__ == '__main__':
    output_dir = r"C:\Users\<user>\Desktop\十五五规划纲要"
    province = sys.argv[1]
    content_file = sys.argv[2]
    with open(content_file, 'r', encoding='utf-8') as f:
        content = f.read()
    create_pdf(province, content, output_dir)
